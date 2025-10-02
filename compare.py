import os
import difflib
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from pypdf import PdfReader
import sys

def getText(fp):
    try:
        if fp.lower().endswith('.pdf'):
            rdr = PdfReader(fp)
            txt = ""
            for pg in rdr.pages:
                txt += pg.extract_text() or ""
            return txt
        elif fp.lower().endswith('.docx'):
            doc = Document(fp)
            txt = "\n".join([p.text for p in doc.paragraphs])
            return txt
        else:
            print(f"Error: bad file type '{fp}'.")
            return None
    except Exception as e:
        print(f"err reading '{fp}': {e}")
        return None

def compareTxt(f1, f2, out):
    print("doing text compare...")
    txt1 = getText(f1)
    txt2 = getText(f2)

    if txt1 is None or txt2 is None:
        print("compare failed, couldn't read a file.")
        return

    m = difflib.SequenceMatcher(a=txt1, b=txt2)
    doc = Document()
    p = doc.add_paragraph()

    RED = RGBColor(255, 0, 0)
    BLUE = RGBColor(0, 0, 255)
    YELLOW = WD_COLOR_INDEX.YELLOW

    def addTxt(chunk, clr=None, hl=None):
        nonlocal p
        paras = chunk.split('\n\n')
        for i, para_txt in enumerate(paras):
            flow = para_txt.replace('\n', ' ')
            
            if flow:
                run = p.add_run(flow)
                if clr:
                    run.font.color.rgb = clr
                if hl:
                    run.font.highlight_color = hl
            
            if i < len(paras) - 1:
                p = doc.add_paragraph()

    for tag, i1, i2, j1, j2 in m.get_opcodes():
        t1 = txt1[i1:i2]
        t2 = txt2[j1:j2]
        
        n1 = ' '.join(t1.split())
        n2 = ' '.join(t2.split())

        if tag == 'equal':
            addTxt(t1)
        elif tag == 'delete':
            addTxt(t1, clr=RED)
        elif tag == 'insert':
            addTxt(t2, clr=BLUE)
        elif tag == 'replace':
            if n1 == n2:
                addTxt(t1, hl=YELLOW)
            else:
                addTxt(t1, clr=RED)
                addTxt(t2, clr=BLUE)
    
    doc.save(out)
    print(f"\ndone! saved to '{os.path.abspath(out)}'")


def cloneP(p_src, doc_tgt, clr=None, hl=None):
    p_new = doc_tgt.add_paragraph()
    p_new.paragraph_format.alignment = p_src.paragraph_format.alignment
    
    for r in p_src.runs:
        r_new = p_new.add_run(r.text)
        r_new.bold = r.bold
        r_new.italic = r.italic
        r_new.underline = r.underline
        r_new.font.name = r.font.name
        r_new.font.size = r.font.size
        if clr:
            r_new.font.color.rgb = clr
        else:
            if r.font.color.rgb:
                r_new.font.color.rgb = r.font.color.rgb
        if hl:
            r_new.font.highlight_color = hl
    return p_new


def compareDocs(f1, f2, out):
    print("doing docx compare...")
    d1 = Document(f1)
    d2 = Document(f2)
    res_doc = Document()

    p1_txt = [p.text for p in d1.paragraphs]
    p2_txt = [p.text for p in d2.paragraphs]

    m = difflib.SequenceMatcher(None, p1_txt, p2_txt)

    RED = RGBColor(255, 0, 0)
    BLUE = RGBColor(0, 0, 255)
    YELLOW = WD_COLOR_INDEX.YELLOW

    for tag, i1, i2, j1, j2 in m.get_opcodes():
        if tag == 'equal':
            for i in range(i1, i2):
                cloneP(d1.paragraphs[i], res_doc)
        
        elif tag == 'insert':
            for i in range(j1, j2):
                cloneP(d2.paragraphs[i], res_doc, clr=BLUE)

        elif tag == 'delete':
            for i in range(i1, i2):
                cloneP(d1.paragraphs[i], res_doc, clr=RED)

        elif tag == 'replace':
            is_space_change = all(
                d1.paragraphs[i1 + i].text.strip() == d2.paragraphs[j1 + i].text.strip()
                for i in range(min(i2 - i1, j2 - j1))
            ) and (i2 - i1 == j2 - j1)

            if is_space_change:
                for i in range(i1, i2):
                    cloneP(d1.paragraphs[i], res_doc, hl=YELLOW)
            else:
                for i in range(i1, i2):
                    cloneP(d1.paragraphs[i], res_doc, clr=RED)
                for i in range(j1, j2):
                    cloneP(d2.paragraphs[i], res_doc, clr=BLUE)
    
    res_doc.save(out)
    print(f"\ndone! saved to '{os.path.abspath(out)}'")


def main():
    print("--- doc compare tool ---")
    
    f1 = input("gimme file 1: ")
    if not os.path.exists(f1):
        print(f"err: can't find '{f1}'.")
        sys.exit(1)

    f2 = input("gimme file 2: ")
    if not os.path.exists(f2):
        print(f"err: can't find '{f2}'.")
        sys.exit(1)

    _, ext1 = os.path.splitext(f1)
    _, ext2 = os.path.splitext(f2)
    if ext1.lower() != ext2.lower():
        print(f"err: files gotta be same type ('{ext1}' and '{ext2}').")
        sys.exit(1)

    b1 = os.path.basename(f1).rsplit('.', 1)[0]
    b2 = os.path.basename(f2).rsplit('.', 1)[0]
    out_file = f"result_{b1}_vs_{b2}.docx"

    if ext1.lower() == '.docx':
        compareDocs(f1, f2, out_file)
    elif ext1.lower() == '.pdf':
        print("\nwarn: comparing pdfs. format will be lost, output is docx.")
        compareTxt(f1, f2, out_file)
    else:
        print(f"err: bad file type '{ext1}'. use docx or pdf.")
        sys.exit(1)

if __name__ == "__main__":
    main()

